import os
import sys
from pathlib import Path

# Add the workspace directory to python path to ensure we can import src
sys.path.append(str(Path(__file__).resolve().parent))

import fitz  # PyMuPDF
import docx
from src.config.settings import UPLOAD_DIR
from src.utils.logger import logger
from src.ingestion.text_parser import PDFParser, DocxParser

def create_sample_files():
    """Generates sample PDF and DOCX files containing English and Tamil text for testing."""
    pdf_path = UPLOAD_DIR / "sample_test.pdf"
    docx_path = UPLOAD_DIR / "sample_test.docx"
    
    logger.info("Generating sample test files...")
    
    # 1. Create a PDF file programmatically using PyMuPDF
    try:
        pdf_doc = fitz.open()
        
        # Page 1: English Text
        page1 = pdf_doc.new_page()
        page1.insert_text((50, 50), "Offline Multimodal RAG Project - Milestone 1", fontsize=16)
        page1.insert_text((50, 100), "This is the first page of our sample PDF document.", fontsize=11)
        page1.insert_text((50, 120), "We are testing PyMuPDF text extraction capabilities.", fontsize=11)
        
        # Page 2: PDF Page 2 Verification (English to avoid default font glyph issue)
        page2 = pdf_doc.new_page()
        page2.insert_text((50, 50), "Milestone 1 - PDF Page 2 Verification", fontsize=16)
        page2.insert_text((50, 100), "This page verifies our multi-page PDF document parsing.", fontsize=12)
        page2.insert_text((50, 125), "Each page is represented as a separate Document object with page metadata.", fontsize=11)
        
        pdf_doc.save(str(pdf_path))
        pdf_doc.close()
        logger.info(f"Sample PDF created at {pdf_path}")
    except Exception as e:
        logger.error(f"Failed to create sample PDF: {e}")
        raise e
        
    # 2. Create a DOCX file programmatically using python-docx
    try:
        docx_doc = docx.Document()
        docx_doc.add_heading("Offline RAG System Verification", level=0)
        
        p1 = docx_doc.add_paragraph("This is a simple paragraph in our DOCX file containing English text.")
        p1.add_run(" Bold text here.").bold = True
        
        p2 = docx_doc.add_paragraph("தமிழ் உரை சோதனை:")
        p2.add_run(" தமிழ் மொழி உலகின் மிக மூத்த மொழிகளில் ஒன்றாகும்.").italic = True
        
        # Add a table
        table = docx_doc.add_table(rows=3, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Language'
        hdr_cells[1].text = 'Greeting'
        
        row1_cells = table.rows[1].cells
        row1_cells[0].text = 'English'
        row1_cells[1].text = 'Hello'
        
        row2_cells = table.rows[2].cells
        row2_cells[0].text = 'Tamil'
        row2_cells[1].text = 'வணக்கம்'
        
        docx_doc.save(str(docx_path))
        logger.info(f"Sample DOCX created at {docx_path}")
    except Exception as e:
        logger.error(f"Failed to create sample DOCX: {e}")
        raise e

    return pdf_path, docx_path

def run_tests():
    """Runs parsers on the generated files and asserts the contents."""
    pdf_path, docx_path = create_sample_files()
    
    # Test PDF Parser
    pdf_parser = PDFParser()
    pdf_docs = pdf_parser.parse(str(pdf_path))
    
    logger.info("=== PDF PARSE RESULTS ===")
    assert len(pdf_docs) == 2, f"Expected 2 pages, got {len(pdf_docs)}"
    
    for i, doc in enumerate(pdf_docs):
        logger.info(f"Page {doc.metadata['page_number']}:")
        logger.info(f"Text Content: {doc.text[:200]}...")
        logger.info(f"Metadata: {doc.metadata}")
        logger.info("-" * 40)
        
    # Verify presence of key text strings
    assert "Offline Multimodal RAG" in pdf_docs[0].text
    assert "PDF Page 2 Verification" in pdf_docs[1].text
    logger.info("PDF extraction assertion check: SUCCESS")
    
    # Test DOCX Parser
    docx_parser = DocxParser()
    docx_docs = docx_parser.parse(str(docx_path))
    
    logger.info("=== DOCX PARSE RESULTS ===")
    # Count should match paragraph/table blocks
    logger.info(f"Total blocks extracted: {len(docx_docs)}")
    
    for doc in docx_docs:
        logger.info(f"Block {doc.metadata['block_index']} ({doc.metadata['block_type']}):")
        logger.info(f"Text Content: {doc.text}")
        logger.info(f"Metadata: {doc.metadata}")
        logger.info("-" * 40)
        
    # Verify table text extraction
    table_docs = [d for d in docx_docs if d.metadata['block_type'] == 'table']
    assert len(table_docs) == 1, "Expected 1 table block"
    assert "Language | Greeting" in table_docs[0].text
    assert "Tamil | வணக்கம்" in table_docs[0].text
    # Verify presence of Tamil paragraph
    assert any("தமிழ் மொழி உலகின் மிக மூத்த" in doc.text for doc in docx_docs), "Tamil paragraph not found in DOCX"
    
    logger.info("DOCX extraction assertion check: SUCCESS")
    logger.info("ALL tests passed successfully!")

if __name__ == "__main__":
    run_tests()
