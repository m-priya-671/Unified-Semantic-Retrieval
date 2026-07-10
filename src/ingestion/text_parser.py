import os
import fitz  # PyMuPDF
import docx
from typing import List
from src.ingestion.base import BaseParser, Document
from src.utils.logger import logger

class PDFParser(BaseParser):
    """Parser to extract text page-by-page from PDF files."""
    
    def parse(self, file_path: str) -> List[Document]:
        """Parses a PDF file and extracts text page-by-page.
        
        Args:
            file_path: Absolute or relative path to the PDF file.
            
        Returns:
            A list of Document objects (one per page).
        """
        documents = []
        logger.info(f"Starting PDF parsing: {file_path}")
        
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
            
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            logger.info(f"Opened PDF {os.path.basename(file_path)} with {total_pages} pages.")
            
            for page_num in range(total_pages):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                # Clean up multiple whitespaces and trim
                clean_text = " ".join(text.split()).strip()
                
                if clean_text:
                    metadata = {
                        "source": os.path.basename(file_path),
                        "file_path": os.path.abspath(file_path),
                        "file_type": "pdf",
                        "page_number": page_num + 1,  # 1-indexed for readability
                        "total_pages": total_pages
                    }
                    documents.append(Document(text=clean_text, metadata=metadata))
                    logger.debug(f"Successfully extracted page {page_num + 1}/{total_pages}")
            
            logger.info(f"Successfully completed PDF parsing: {file_path}. Extracted {len(documents)} pages.")
        except Exception as e:
            logger.error(f"Error occurred while parsing PDF {file_path}: {str(e)}")
            raise e
            
        return documents

class DocxParser(BaseParser):
    """Parser to extract text sequentially (paragraphs & tables) from Word documents (DOCX)."""
    
    def parse(self, file_path: str) -> List[Document]:
        """Parses a DOCX file and extracts paragraphs and tables in sequential reading order.
        
        Args:
            file_path: Absolute or relative path to the DOCX file.
            
        Returns:
            A list of Document objects (one per block item).
        """
        documents = []
        logger.info(f"Starting DOCX parsing: {file_path}")
        
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
            
        try:
            doc = docx.Document(file_path)
            
            # Helper function to extract text from a table and format it cleanly
            def extract_table_text(table) -> str:
                table_lines = []
                for row in table.rows:
                    # Clean and extract text from each cell in the row
                    cells_text = []
                    for cell in row.cells:
                        cell_txt = " ".join(cell.text.split()).strip()
                        cells_text.append(cell_txt)
                    
                    # Joining cells with markdown-like pipeline character
                    row_line = " | ".join(filter(None, cells_text))
                    if row_line:
                        table_lines.append(row_line)
                return "\n".join(table_lines)
            
            body_elements = []
            
            # Iterate through body elements to maintain exact reading order
            for child in doc.element.body:
                if child.tag.endswith('p'):
                    p = docx.text.paragraph.Paragraph(child, doc)
                    text = " ".join(p.text.split()).strip()
                    if text:
                        body_elements.append(("paragraph", text))
                elif child.tag.endswith('tbl'):
                    t = docx.table.Table(child, doc)
                    text = extract_table_text(t)
                    if text:
                        body_elements.append(("table", text))
            
            # Package blocks as Document objects
            total_elements = len(body_elements)
            for index, (block_type, text) in enumerate(body_elements):
                metadata = {
                    "source": os.path.basename(file_path),
                    "file_path": os.path.abspath(file_path),
                    "file_type": "docx",
                    "block_type": block_type,
                    "block_index": index + 1,  # 1-indexed
                    "total_blocks": total_elements
                }
                documents.append(Document(text=text, metadata=metadata))
            
            logger.info(f"Successfully completed DOCX parsing: {file_path}. Extracted {len(documents)} blocks.")
        except Exception as e:
            logger.error(f"Error occurred while parsing DOCX {file_path}: {str(e)}")
            raise e
            
        return documents
